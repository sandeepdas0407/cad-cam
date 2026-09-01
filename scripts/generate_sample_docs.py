"""Generates a sample_docs/ folder with test files for manual + e2e verification:
- a text-based PDF (real selectable text)
- a scanned/image-only PDF (no text layer, text baked into a rendered image)
- a mixed PDF (real text title block + an embedded image with baked-in text)
- a Word document with a heading and a table
- an unsupported CAD file and an image file, to confirm the walker skips non-PDF/DOCX files
"""

import io
from pathlib import Path

import docx
import fitz
from PIL import Image, ImageDraw, ImageFont

OUT = Path(__file__).resolve().parent.parent / "sample_docs"
OUT.mkdir(exist_ok=True)


def _font(size=28):
    try:
        return ImageFont.truetype("arial.ttf", size)
    except Exception:
        return ImageFont.load_default()


def make_text_image(lines: list[str], size=(1000, 700)) -> Image.Image:
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    y = 30
    for line in lines:
        draw.text((30, y), line, fill="black", font=_font(26))
        y += 40
    draw.rectangle([10, 10, size[0] - 10, size[1] - 10], outline="black", width=3)
    return img


def make_text_pdf():
    path = OUT / "motor_mount_spec.pdf"
    doc = fitz.open()
    page = doc.new_page()
    text = (
        "MOTOR MOUNT ASSEMBLY - ENGINEERING SPECIFICATION\n\n"
        "Part Number: MM-4471-B\n"
        "Material: 6061-T6 Aluminum Alloy\n\n"
        "Torque Specification:\n"
        "All M8 mounting bolts shall be torqued to 24 Nm +/- 2 Nm using a "
        "calibrated torque wrench, applied in a star pattern sequence.\n\n"
        "Tolerance:\n"
        "Bore diameter tolerance is H7. Surface finish on the mounting face "
        "shall not exceed Ra 1.6 micrometers.\n\n"
        "Fastener retention shall be verified using thread-locking compound "
        "Loctite 243 on all mounting bolts.\n"
    )
    page.insert_text((50, 60), text, fontsize=13, fontname="helv")
    doc.save(str(path))
    doc.close()
    print(f"wrote {path}")


def make_scanned_pdf():
    path = OUT / "shaft_drawing_scanned.pdf"
    img = make_text_image(
        [
            "DRIVE SHAFT - SHEET 1 OF 1",
            "PART NO: DS-2209-A",
            "",
            "SHAFT DIAMETER: 25.4 mm h6",
            "KEYWAY WIDTH: 6.0 mm +0.02/-0.00",
            "OVERALL LENGTH: 340 mm",
            "",
            "SPLINE FIT TOLERANCE CLASS: 5",
            "SURFACE TREATMENT: BLACK OXIDE",
        ]
    )
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    doc = fitz.open()
    page = doc.new_page()
    rect = fitz.Rect(50, 50, 50 + img.width * 0.5, 50 + img.height * 0.5)
    page.insert_image(rect, stream=buf.getvalue())
    doc.save(str(path))
    doc.close()
    print(f"wrote {path}")


def make_mixed_pdf():
    path = OUT / "bracket_drawing_mixed.pdf"
    img = make_text_image(
        [
            "MOUNTING BRACKET DETAIL",
            "",
            "HOLE PATTERN: 4x M6 CLEARANCE",
            "BEND RADIUS: 3.0 mm MINIMUM",
            "SHEET THICKNESS: 2.5 mm COLD ROLLED STEEL",
            "FLATNESS TOLERANCE: 0.1 mm ACROSS FULL LENGTH",
        ],
        size=(900, 500),
    )
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    doc = fitz.open()
    page = doc.new_page()
    title_block_text = (
        "TITLE BLOCK\n"
        "Drawing No: BR-1187-C   Rev: B\n"
        "Drawn by: J. Alvarez   Approved: R. Chen\n"
        "Scale: 1:2   Units: mm\n"
        "Anodized finish per spec AN-200, clear coat.\n"
    )
    page.insert_text((50, 40), title_block_text, fontsize=12, fontname="helv")
    rect = fitz.Rect(50, 180, 50 + img.width * 0.5, 180 + img.height * 0.5)
    page.insert_image(rect, stream=buf.getvalue())
    doc.save(str(path))
    doc.close()
    print(f"wrote {path}")


def make_docx():
    path = OUT / "assembly_notes.docx"
    document = docx.Document()
    document.add_heading("Gearbox Assembly Notes", level=1)
    document.add_paragraph(
        "This document covers assembly notes and fastener tolerances for the "
        "GB-300 series gearbox housing."
    )
    document.add_heading("Torque Tolerances", level=2)
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Fastener", "Torque (Nm)", "Tolerance"
    rows = [
        ("M6 Housing Bolt", "9", "+/- 1 Nm"),
        ("M10 Shaft Retainer", "45", "+/- 3 Nm"),
        ("M4 Cover Screw", "2.5", "+/- 0.3 Nm"),
    ]
    for f, t, tol in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = f, t, tol

    document.add_heading("Lubrication", level=2)
    document.add_paragraph(
        "Apply synthetic gear oil ISO VG 220 to all meshing gear surfaces prior "
        "to final housing closure. Verify oil level using the sight glass on the "
        "lower housing before commissioning."
    )
    document.save(str(path))
    print(f"wrote {path}")


def make_ignored_files():
    (OUT / "gearbox_model.dxf").write_text(
        "0\nSECTION\n2\nHEADER\n0\nENDSEC\n0\nEOF\n", encoding="utf-8"
    )
    img = make_text_image(["REFERENCE PHOTO - NOT INDEXED"], size=(400, 200))
    img.save(OUT / "reference_photo.jpg")
    print(f"wrote {OUT / 'gearbox_model.dxf'} and reference_photo.jpg (should be skipped by the walker)")


if __name__ == "__main__":
    make_text_pdf()
    make_scanned_pdf()
    make_mixed_pdf()
    make_docx()
    make_ignored_files()
    print("Done.")
