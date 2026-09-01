import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import ParsedDocument, ParsedPage

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"}


def _iter_body_items(document):
    """Yields Paragraph/Table objects in true document order (document.paragraphs
    and document.tables are separate flat lists that lose interleaving order)."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def parse_docx(path: str) -> ParsedDocument:
    doc = ParsedDocument(path=path, doc_type="docx")
    document = docx.Document(path)

    current_heading = None
    buffer: list[str] = []

    def flush():
        text = "\n".join(buffer).strip()
        if text:
            doc.pages.append(
                ParsedPage(
                    page_number=None,
                    text=text,
                    source="text_layer",
                    section_heading=current_heading,
                )
            )
        buffer.clear()

    for item in _iter_body_items(document):
        if isinstance(item, Paragraph):
            style_name = item.style.name if item.style else ""
            if style_name in HEADING_STYLES:
                flush()
                current_heading = item.text.strip() or current_heading
                continue
            if item.text.strip():
                buffer.append(item.text)
        elif isinstance(item, Table):
            flush()
            rows_text = []
            for row in item.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            table_text = "\n".join(rows_text).strip()
            if table_text:
                doc.pages.append(
                    ParsedPage(
                        page_number=None,
                        text=table_text,
                        source="text_layer",
                        section_heading=current_heading,
                    )
                )

    flush()
    return doc
